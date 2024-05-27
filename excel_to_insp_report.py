import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
import pandas as pd
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
from oauth2client.service_account import ServiceAccountCredentials
import gspread

import subprocess

import os

# Function to download the image from a URL
def download_image(service_account_file, file_path, save_path):
    # Authenticate using service account credentials
    credentials = service_account.Credentials.from_service_account_file(service_account_file)
    drive_service = build('drive', 'v3', credentials=credentials)
    file_id = get_file_id_from_url(file_path)
    # Download the image
    request = drive_service.files().get_media(fileId=file_id)
    fh = open(save_path, "wb")
    downloader = MediaIoBaseDownload(fh, request)
    done = False
    while done is False:
        status, done = downloader.next_chunk()
        print("Download %d%%." % int(status.progress() * 100))
    print("Image downloaded successfully.")

def get_file_id_from_url(url):
    # Extract the file ID from the URL
    return url.split('/')[-2]

############################ Google Sheets ##########################################
import gspread

# Define the scope and credentials for Google Sheets API
scope = ['https://spreadsheets.google.com/feeds',
   'https://www.googleapis.com/auth/drive']

# Authorize the client with credentials
creds = ServiceAccountCredentials.from_json_keyfile_name('credentials.json', scope)
client = gspread.authorize(creds)

# Open the Google Spreadsheet
spreadsheet_name = 'Safety-Records'
inspection_sheet = client.open(spreadsheet_name).worksheet('Inspection')
meeting_sheet = client.open(spreadsheet_name).worksheet('Meeting')
training_sheet = client.open(spreadsheet_name).worksheet('Training')

# Get all records in the Google Spreadsheet
inspection_data = inspection_sheet.get_all_records()
meeting_data = meeting_sheet.get_all_records()
training_data = training_sheet.get_all_records()

# Load data into a Pandas DataFrame
inspection_df = pd.DataFrame(inspection_data)
meeting_df = pd.DataFrame(meeting_data)
training_df = pd.DataFrame(meeting_data)
######################################################################################


#date = "23-05-2024"
#department = "BF"

def create_report(date, department, file_path):
    doc = docx.Document()

    para=doc.add_paragraph()
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r=para.add_run(department)
    r.bold = True

    para=doc.add_paragraph()
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r=para.add_run('Inspection Observations')
    r.bold = True

    para=doc.add_paragraph()
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r=para.add_run(date)
    r.bold = True

    para=doc.add_paragraph()

    observations_df = pd.DataFrame(inspection_df.loc[(inspection_df['Inspection Date']==date) & (inspection_df['Department']==department)])

    table = doc.add_table(rows=observations_df.shape[0]+1, cols=4)
    table.style = 'Table Grid'
    row = table.rows[0].cells
    row[0].text = "Location"
    row[0].paragraphs[0].runs[0].font.bold = True
    row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    row[1].text = "Observation"
    row[1].paragraphs[0].runs[0].font.bold = True
    row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    row[2].text = "Picture 1"
    row[2].paragraphs[0].runs[0].font.bold = True
    row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    row[3].text = "Picture 2"
    row[3].paragraphs[0].runs[0].font.bold = True
    row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    i=1
    for index, observation in observations_df.iterrows():
        row = table.rows[i].cells
        row[0].text = observation['Location']
        row[1].text = observation['Observation']

        if (len(observation['Photo'].split(","))==1):
            paragraph = row[2].paragraphs[0]
            image_1 = observation['Photo'].split(",")[0]  #get a list of links to all images of the particular observation
            download_image( 'credentials.json', image_1, 'image.jpg')
            paragraph.add_run().add_picture('image.jpg', width=Inches(2.0), height=Inches(2.0))
            # Remove the image from local system
            os.remove('image.jpg')
        elif (len(observation['Photo'].split(","))>1):
            paragraph = row[2].paragraphs[0]
            image_1 = observation['Photo'].split(",")[0]  #get a list of links to all images of the particular observation
            download_image('credentials.json', image_1, 'image.jpg')
            paragraph.add_run().add_picture('image.jpg', width=Inches(2.0), height=Inches(2.0))
            # Remove the image from local system
            os.remove('image.jpg')

            paragraph = row[3].paragraphs[0]
            image_2 = observation['Photo'].split(",")[1]  #get a list of links to all images of the particular observation
            download_image( 'credentials.json', image_2, 'image.jpg')
            paragraph.add_run().add_picture('image.jpg', width=Inches(2.0), height=Inches(2.0))
            # Remove the image from local system
            os.remove('image.jpg')
        else:
            pass
        i+=1

    # Save the document as a Word file
    doc.save(file_path)

    #print(f"PDF file saved: {pdf_file}")

    #doc.save(r'C:\Users\709412\Downloads\test.docx')